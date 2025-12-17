import docx
from pypdf import PdfReader

def read_docx(path):
    print(f"\n--- Reading DOCX: {path} ---")
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        print(f"Success! Read {len(doc.paragraphs)} paragraphs.")
        print(f"snippet: {' '.join(full_text)[:300]}...")
        return True
    except Exception as e:
        print(f"Failed to read DOCX: {e}")
        return False

def read_pdf(path):
    print(f"\n--- Reading PDF: {path} ---")
    try:
        reader = PdfReader(path)
        print(f"Success! PDF has {len(reader.pages)} pages.")
        first_page = reader.pages[0].extract_text()
        print(f"snippet: {first_page[:300]}...")
        return True
    except Exception as e:
        print(f"Failed to read PDF: {e}")
        return False

if __name__ == "__main__":
    d = read_docx("competition.docx")
    p = read_pdf("case_study.pdf")
    
    if d and p:
        print("\nOVERALL STATUS: SUCCESS")
    else:
        print("\nOVERALL STATUS: PARTIAL FAILURE")
